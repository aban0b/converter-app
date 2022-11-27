from PIL import Image
import pytesseract
from pydub import AudioSegment
from googletrans import Translator
import subprocess
import codecs
from docx import Document
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
import sys
from PyQt5.uic import loadUiType
import time
from os import path
ui,_ = loadUiType('main.ui')
class MainApp1(QMainWindow , ui):
    c = ['english', 'afrikaans', 'albanian', 'amharic', 'arabic', 'armenian', 'azerbaijani', 'basque', 'belarusian',
         'bengali',
         'bosnian', 'bulgarian', 'catalan', 'cebuano', 'chichewa', 'chinese (simplified)', 'chinese (traditional)',
         'corsican', 'croatian', 'czech', 'danish', 'dutch', 'esperanto', 'estonian', 'filipino',
         'finnish',
         'french', 'frisian', 'galician', 'georgian', 'german', 'greek', 'gujarati', 'haitian creole', 'hausa',
         'hawaiian', 'hebrew', 'hindi', 'hmong', 'hungarian', 'icelandic', 'igbo', 'indonesian', 'irish', 'italian',
         'japanese', 'javanese', 'kannada', 'kazakh', 'khmer', 'korean', 'kurdish (kurmanji)', 'kyrgyz', 'lao',
         'latin',
         'latvian', 'lithuanian', 'luxembourgish', 'macedonian', 'malagasy', 'malay', 'malayalam', 'maltese',
         'maori',
         'marathi', 'mongolian', 'myanmar (burmese)', 'nepali', 'norwegian', 'pashto', 'persian', 'polish',
         'portuguese', 'punjabi', 'romanian', 'russian', 'samoan', 'scots gaelic', 'serbian', 'sesotho', 'shona',
         'sindhi', 'sinhala', 'slovak', 'slovenian', 'somali', 'spanish', 'sundanese', 'swahili', 'swedish',
         'tajik',
         'tamil', 'telugu', 'thai', 'turkish', 'ukrainian', 'urdu', 'uzbek', 'vietnamese', 'welsh', 'xhosa',
         'yiddish',
         'yoruba', 'zulu', 'Filipino', 'Hebrew']
    def __init__(self , parent=None):
        super(MainApp1 , self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        for i in self.c:
            self.comboBox.addItem(i)
        self.InitUI()
        self.Handel_Buttons()

    def Handel_Buttons(self):
            self.pushButton_8.clicked.connect(self.tab1)
            self.pushButton_9.clicked.connect(self.tab2)
            self.pushButton_11.clicked.connect(self.tab3)
            self.pushButton_10.clicked.connect(self.tab4)
            self.pushButton_12.clicked.connect(self.tab0)
            self.pushButton.clicked.connect(self.confert_png_to_ico)
            self.pushButton_2.clicked.connect(self.confert_jpg_to_ico)
            self.pushButton_3.clicked.connect(self.confert_ico_to_png)
            self.pushButton_4.clicked.connect(self.confert_png_to_jpg)
            self.pushButton_5.clicked.connect(self.confert_jpg_to_png)
            self.pushButton_6.clicked.connect(self.get_file_text)
            self.pushButton_7.clicked.connect(self.get_word_text)
            self.pushButton_14.clicked.connect(self.write_to_word)
            self.pushButton_15.clicked.connect(self.write_to_file)
            self.pushButton_16.clicked.connect(self.confert_mp4_to_mp3)
            self.pushButton_37.clicked.connect(self.confert_wav_to_mp3)
            self.pushButton_38.clicked.connect(self.confert_mp3_to_wav)
            self.pushButton_13.clicked.connect(self.get_photo_txt)
            self.pushButton_17.clicked.connect(self.get_photo_txt1)
            self.pushButton_20.clicked.connect(self.trans_fun)

    def trans_fun(self):
        response0 = str(self.textEdit_3.toPlainText())
        response1 = " to "
        trans_to = self.comboBox.currentIndex()
        response2 = str(self.c[trans_to])
        translator = Translator()  # Create object of Translator.
        text = response0 + response1 + response2
        text = text.split()
        n = len(text)
        lang = text[n - 1]
        text[n - 2] = ''
        text[n - 1] = ''
        data = ' '.join(text)
        translated = translator.translate(data, dest='%s' % lang)
        data = (translated.text)
        self.textEdit_3.setText(str(data))
    def get_photo_txt1(self):
        try:
            url_dir = QFileDialog.getOpenFileName(self, 'Select images', '', 'photo (*)')
            name = url_dir[0]
            img = Image.open(name)
            pytesseract.pytesseract.tesseract_cmd = 'C:\Program Files\Tesseract-OCR/tesseract.exe'
            result = pytesseract.image_to_string(img,lang='ara')
            self.textEdit_3.setText(result)
        except :
            pass
    def get_photo_txt(self):
        try:
            url_dir = QFileDialog.getOpenFileName(self, 'Select images', '', 'photo (*)')
            name = url_dir[0]
            img = Image.open(name)
            pytesseract.pytesseract.tesseract_cmd = 'C:\Program Files\Tesseract-OCR/tesseract.exe'
            result = pytesseract.image_to_string(img)
            self.textEdit_3.setText(result)
        except :
            pass
    def confert_mp3_to_wav(self):
      try :
        url_dir = QFileDialog.getOpenFileName(self, 'Select videos', '', 'videos (*.mp3)')
        name = url_dir[0]
        hisname = path.splitext(name)[0]
        src = name
        dst = "{}.wav".format(hisname)
        # convert wav to mp3
        sound = AudioSegment.from_mp3(src)
        sound.export(dst, format="wav")
        QMessageBox.about(self, 'covert', 'you convert it know')
      except :
         QMessageBox.about(self, 'error', 'you have error try again')
    def confert_wav_to_mp3(self):
      try :
        url_dir = QFileDialog.getOpenFileName(self, 'Select videos', '', 'videos (*.wav)')
        name = url_dir[0]
        hisname = path.splitext(name)[0]
        sound = AudioSegment.from_wav(name)
        sound.export("{}.mp3".format(hisname), format="mp3")
        QMessageBox.about(self, 'covert', 'you convert it know')
      except:
        QMessageBox.about(self, 'error', 'you have error try again')
    def confert_mp4_to_mp3(self):
     try:
        url_dir = QFileDialog.getOpenFileName(self, 'Select videos', '', 'videos (*.mp4)')
        name = url_dir[0]
        hisname = path.splitext(name)[0]
        command = "ffmpeg -i {} -ab 160k -ac 2 -ar 44100 -vn {}.mp3".format(name, hisname)
        subprocess.call(command, shell=True)
        QMessageBox.about(self, 'covert', 'you convert it know')
     except:
        QMessageBox.about(self, 'error', 'you have error try again')
    def write_to_file(self):
      try :
        url_output = QFileDialog.getSaveFileName(self, 'Save As', '', 'text Files (*.txt)')
        text = self.textEdit_2.toPlainText()
        f = open(url_output[0], "w", encoding='utf-8')
        f.write(text)
        f.close()
        QMessageBox.about(self, 'covert', 'you convert it know')
      except :
          QMessageBox.about(self, 'error', 'you have error try again')

    def get_word_text(self):
     try:
        url_dir = QFileDialog.getOpenFileName(self, 'Select text files', '', 'Text Files (*.docx)')
        name = url_dir[0]
        doc = Document(name)
        c = []
        for i in doc.paragraphs:
            c.append(i.text)
        r = '\n'.join(c)
        self.label_32.setText("سيتم وضع هذا النص فى ملف ورد يمكنك التعديل عليه")
        self.textEdit_2.setText(r)
     except :
         QMessageBox.about(self, 'error', 'you forget choose word file')

    def write_to_word(self):
      try :
        url_output = QFileDialog.getSaveFileName(self, 'Save As', '', 'Word Files (*.docx)')
        text=self.textEdit.toPlainText()
        document = Document()
        p1 = document.add_paragraph(text)
        document.save(r"{}".format(url_output[0]))
        QMessageBox.about(self, 'covert', 'you convert it know')
      except :
          QMessageBox.about(self, 'error', 'you have error try again')
    def get_file_text(self):
      try :
        url_dir = QFileDialog.getOpenFileName(self, 'Select text files', '', 'Text Files (*.txt)')
        name = url_dir[0]
        f = codecs.open(name, encoding='utf-8')
        c = ""
        for line in f:
            c += line
        c = c.split()
        c = " ".join(c)
        self.label_31.setText("سيتم وضع هذا النص فى ملف ورد يمكنك التعديل عليه")
        self.textEdit.setText(c)
      except:
          QMessageBox.about(self, 'error', 'you forget choose file text')
    def confert_jpg_to_png(self):
        try:
          url_dir = QFileDialog.getOpenFileName(self, 'Select image', '', 'Image Files (*.jpg)')
          name = url_dir[0]
          im1 = Image.open(r'{}'.format(name))
          hisname = path.splitext(name)[0]
          im1.save(r'{}.png'.format(hisname))
          QMessageBox.about(self, 'done', 'convert sucess !')
        except:
            QMessageBox.about(self, 'error', 'you forget choose picture')

    def confert_png_to_jpg(self):
        try:
         url_dir = QFileDialog.getOpenFileName(self, 'Select image', '', 'Image Files (*.png)')
         name = url_dir[0]
         im1 = Image.open(r'{}'.format(name))
         hisname = path.splitext(name)[0]
         im1.save(r'{}.jpg'.format(hisname))
         QMessageBox.about(self, 'done', 'convert sucess !')
        except:
            QMessageBox.about(self, 'error', 'you forget choose picture')
    def confert_ico_to_png(self):
        try :
            url_dir = QFileDialog.getOpenFileName(self, 'Select image', '', 'Image Files (*.ico)')
            name = url_dir[0]
            hisname = path.splitext(name)[0]
            filename = r'{}'.format(name)
            img = Image.open(filename)
            img.save('{}.png'.format(hisname))
            QMessageBox.about(self, 'done', 'convert sucess !')
        except :
            QMessageBox.about(self, 'error', 'you forget choose picture')
    def confert_png_to_ico(self):
        try:
            url_dir = QFileDialog.getOpenFileName(self, 'Select image', '', 'Image Files (*.png)')
            name = url_dir[0]
            hisname = path.splitext(name)[0]
            filename = r'{}'.format(name)
            img = Image.open(filename)
            img.save('{}.ico'.format(hisname))
            QMessageBox.about(self, 'done', 'convert sucess !')
        except :
            QMessageBox.about(self, 'error', 'you forget choose picture')
    def confert_jpg_to_ico(self):
        try:
            url_dir = QFileDialog.getOpenFileName(self, 'Select image', '', 'Image Files (*.jpg)')
            name = url_dir[0]
            hisname = path.splitext(name)[0]
            filename = r'{}'.format(name)
            img = Image.open(filename)
            img.save('{}.ico'.format(hisname))
            QMessageBox.about(self, 'done', 'convert sucess !')
        except :
            QMessageBox.about(self, 'error', 'you forget choose picture')
    def tab0(self):
            self.tabWidget.setCurrentIndex(0)
            self.InitUI()

    def tab1(self):
            self.tabWidget.setCurrentIndex(1)
            self.InitUI1()

    def tab2(self):
            self.tabWidget.setCurrentIndex(2)
            self.InitUI2()

    def tab3(self):
            self.tabWidget.setCurrentIndex(3)
            self.InitUI3()
    def tab4(self):
            self.tabWidget.setCurrentIndex(4)
            self.InitUI4()

    def InitUI(self):
        ## contain all ui changes in loading
        #self.tabWidget.tabBar().setVisible(False)
        self.Move_Box_1()
        self.Move_Box_2()
        self.Move_Box_3()
        self.Move_Box_4()
        self.Move_Box_5()
    def InitUI1(self):
        self.Move_Box_6()
        self.Move_Box_7()
        self.Move_Box_8()
        self.Move_Box_9()
        self.Move_Box_10()
    def InitUI2(self):
        self.Move_Box_11()
        self.Move_Box_12()
        self.Move_Box_133()
    def InitUI3(self):
        self.Move_Box_13()
        self.Move_Box_14()
        self.Move_Box_15()
    def InitUI4(self):
        self.Move_Box_16()
        self.Move_Box_17()
    def Move_Box_1(self):
        box_animation1 = QPropertyAnimation(self.label_2 , b"geometry")
        box_animation1.setDuration(2500)
        box_animation1.setStartValue(QRect(0,0,0,0))
        box_animation1.setEndValue(QRect(20,40,71,71))
        box_animation1.start()
        self.box_animation1 = box_animation1
    def Move_Box_2(self):
        box_animation2 = QPropertyAnimation(self.label_3 , b"geometry")
        box_animation2.setDuration(2500)
        box_animation2.setStartValue(QRect(0,0,0,0))
        box_animation2.setEndValue(QRect(30,30,61,71))
        box_animation2.start()
        self.box_animation2 = box_animation2


    def Move_Box_3(self):
        box_animation3 = QPropertyAnimation(self.label_4 , b"geometry")
        box_animation3.setDuration(2500)
        box_animation3.setStartValue(QRect(0,0,0,0))
        box_animation3.setEndValue(QRect(10,40,81,61))
        box_animation3.start()
        self.box_animation3 = box_animation3


    def Move_Box_4(self):
        box_animation4 = QPropertyAnimation(self.label_5 , b"geometry")
        box_animation4.setDuration(2500)
        box_animation4.setStartValue(QRect(0,0,0,0))
        box_animation4.setEndValue(QRect(26,30,61,51))
        box_animation4.start()
        self.box_animation4 = box_animation4
    def Move_Box_5(self):
        box_animation5 = QPropertyAnimation(self.label , b"geometry")
        box_animation5.setDuration(2500)
        box_animation5.setStartValue(QRect(0,0,0,0))
        box_animation5.setEndValue(QRect(480,160,91,81))
        box_animation5.start()
        self.box_animation5 = box_animation5
    def Move_Box_6(self):
        box_animation6 = QPropertyAnimation(self.label_28 , b"geometry")
        box_animation6.setDuration(2500)
        box_animation6.setStartValue(QRect(0,0,0,0))
        box_animation6.setEndValue(QRect(90,0,491,421))
        box_animation6.start()
        self.box_animation6 = box_animation6
    def Move_Box_7(self):
        box_animation7 = QPropertyAnimation(self.groupBox_6 , b"geometry")
        box_animation7.setDuration(2500)
        box_animation7.setStartValue(QRect(0,0,0,0))
        box_animation7.setEndValue(QRect(680,10,311,131))
        box_animation7.start()
        self.box_animation7 = box_animation7
    def Move_Box_8(self):
        box_animation8 = QPropertyAnimation(self.groupBox_7 , b"geometry")
        box_animation8.setDuration(2500)
        box_animation8.setStartValue(QRect(0,0,0,0))
        box_animation8.setEndValue(QRect(680,150,311,91))
        box_animation8.start()
        self.box_animation8 = box_animation8
    def Move_Box_9(self):
        box_animation9 = QPropertyAnimation(self.groupBox_8 , b"geometry")
        box_animation9.setDuration(2500)
        box_animation9.setStartValue(QRect(0,0,0,0))
        box_animation9.setEndValue(QRect(680,250,311,91))
        box_animation9.start()
        self.box_animation9 = box_animation9
    def Move_Box_10(self):
        box_animation10 = QPropertyAnimation(self.groupBox_9 , b"geometry")
        box_animation10.setDuration(2500)
        box_animation10.setStartValue(QRect(0,0,0,0))
        box_animation10.setEndValue(QRect(680,330,311,91))
        box_animation10.start()
        self.box_animation10 = box_animation10
    def Move_Box_11(self):
        box_animation11 = QPropertyAnimation(self.label_106 , b"geometry")
        box_animation11.setDuration(2500)
        box_animation11.setStartValue(QRect(0,0,0,0))
        box_animation11.setEndValue(QRect(40,10,521,411))
        box_animation11.start()
        self.box_animation11 = box_animation11
    def Move_Box_12(self):
        box_animation12 = QPropertyAnimation(self.groupBox_12 , b"geometry")
        box_animation12.setDuration(2500)
        box_animation12.setStartValue(QRect(0,0,0,0))
        box_animation12.setEndValue(QRect(720,30,271,101))
        box_animation12.start()
        self.box_animation12 = box_animation12
    def Move_Box_133(self):
        box_animation133 = QPropertyAnimation(self.groupBox_13 , b"geometry")
        box_animation133.setDuration(2500)
        box_animation133.setStartValue(QRect(0,0,0,0))
        box_animation133.setEndValue(QRect(720,200,271,121))
        box_animation133.start()
        self.box_animation133 = box_animation133
    def Move_Box_13(self):
        box_animation13 = QPropertyAnimation(self.groupBox_11 , b"geometry")
        box_animation13.setDuration(2500)
        box_animation13.setStartValue(QRect(0,0,0,0))
        box_animation13.setEndValue(QRect(60,30,381,341))
        box_animation13.start()
        self.box_animation13 = box_animation13
    def Move_Box_14(self):
        box_animation14 = QPropertyAnimation(self.groupBox_10 , b"geometry")
        box_animation14.setDuration(2500)
        box_animation14.setStartValue(QRect(0,0,0,0))
        box_animation14.setEndValue(QRect(610,30,371,341))
        box_animation14.start()
        self.box_animation14 = box_animation14
    def Move_Box_15(self):
        box_animation15 = QPropertyAnimation(self.label_34 , b"geometry")
        box_animation15.setDuration(2500)
        box_animation15.setStartValue(QRect(0,0,0,0))
        box_animation15.setEndValue(QRect(490,160,81,71))
        box_animation15.start()
        self.box_animation15 = box_animation15
    def Move_Box_16(self):
        box_animation16 = QPropertyAnimation(self.label_22 , b"geometry")
        box_animation16.setDuration(2500)
        box_animation16.setStartValue(QRect(0,0,0,0))
        box_animation16.setEndValue(QRect(-40,-40,341,421))
        box_animation16.start()
        self.box_animation16 = box_animation16
    def Move_Box_17(self):
        box_animation17 = QPropertyAnimation(self.label_6 , b"geometry")
        box_animation17.setDuration(2500)
        box_animation17.setStartValue(QRect(0,0,0,0))
        box_animation17.setEndValue(QRect(240,10,741,51))
        box_animation17.start()
        self.box_animation17 = box_animation17
        box_animation18 = QPropertyAnimation(self.label_7, b"geometry")
        box_animation18.setDuration(2500)
        box_animation18.setStartValue(QRect(0, 0, 0, 0))
        box_animation18.setEndValue(QRect(240, 70, 741, 51))
        box_animation18.start()
        self.box_animation18 = box_animation18
        box_animation19 = QPropertyAnimation(self.label_16, b"geometry")
        box_animation19.setDuration(2500)
        box_animation19.setStartValue(QRect(0, 0, 0, 0))
        box_animation19.setEndValue(QRect(240, 130, 741, 51))
        box_animation19.start()
        self.box_animation19 = box_animation19
        box_animation20 = QPropertyAnimation(self.label_17, b"geometry")
        box_animation20.setDuration(2500)
        box_animation20.setStartValue(QRect(0, 0, 0, 0))
        box_animation20.setEndValue(QRect(240, 190, 741, 51))
        box_animation20.start()
        self.box_animation20 = box_animation20
        box_animation21 = QPropertyAnimation(self.label_18, b"geometry")
        box_animation21.setDuration(2500)
        box_animation21.setStartValue(QRect(0, 0, 0, 0))
        box_animation21.setEndValue(QRect(240, 250, 741, 51))
        box_animation21.start()
        self.box_animation21 = box_animation21
        box_animation22 = QPropertyAnimation(self.label_19, b"geometry")
        box_animation22.setDuration(2500)
        box_animation22.setStartValue(QRect(0, 0, 0, 0))
        box_animation22.setEndValue(QRect(240, 310, 741, 51))
        box_animation22.start()
        self.box_animation22 = box_animation22
        box_animation23 = QPropertyAnimation(self.label_20, b"geometry")
        box_animation23.setDuration(2500)
        box_animation23.setStartValue(QRect(0, 0, 0, 0))
        box_animation23.setEndValue(QRect(240, 370, 741, 51))
        box_animation23.start()
        self.box_animation23 = box_animation23



LIMIT_TIME = 100
FORM_CLASS, _ = loadUiType(path.join(path.dirname(__file__), 'first_page.ui'))
class Ex(QThread):
    countChanged = pyqtSignal(int)
    def run(self):
        count = 0
        while count < LIMIT_TIME:
            count += 1
            time.sleep(0.04)
            self.countChanged.emit(count)
class mainApp(QMainWindow, FORM_CLASS):
    def __init__(self, parent=None):
        super(mainApp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.progressBar.setValue(0)
        self.onStart()
    def onStart(self):
        self.calc = Ex()
        self.calc.countChanged.connect(self.countChangedProcess)
        self.calc.start()
    def countChangedProcess(self, value):
        self.progressBar.setValue(value)
        if value == 100:
            QApplication.processEvents()
            time.sleep(0.04)
            self.closeCurrentApp_OpenNewApp()
    def closeCurrentApp_OpenNewApp(self):
        self.close()
        self.open = MainApp1()
        self.open.show()
app = QApplication(sys.argv)
window = mainApp()
window.show()
app.exec_()
